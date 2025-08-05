import os
import time
import logging
import hashlib
from datetime import datetime
from typing import List, Optional, Dict, Any
from contextlib import asynccontextmanager
import requests
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document as DocxDocument
import email
import io
import numpy as np
import asyncio
from concurrent.futures import ThreadPoolExecutor
import aiohttp
from functools import lru_cache

from fastapi import FastAPI, HTTPException, Header, Depends, BackgroundTasks
from pydantic import BaseModel

# RAG imports
from langchain_together import ChatTogether, TogetherEmbeddings
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

EXPECTED_TOKEN = "5aa05ad358e859e92978582cde20423149f28beb49da7a2bbb487afa8fce1be8"

class QuestionRequest(BaseModel):
    documents: str
    questions: List[str]

class AnswerResponse(BaseModel):
    answers: List[str]

class UltraSpeedVectorStore:
    def __init__(self, embeddings):
        self.embeddings = embeddings
        self.documents = []
        self.vectors = []
        self._embedding_cache = {}
    
    async def add_documents_ultra_fast(self, documents: List[Document]):
        """ULTRA SPEED OPTIMIZED - Batch embeddings"""
        logger.info(f"ULTRA SPEED: Processing {len(documents)} chunks")
        
        start_time = time.time()
        
        # Extract unique texts for batch embedding
        texts = []
        text_to_indices = {}
        
        for i, doc in enumerate(documents):
            text = doc.page_content
            text_hash = hashlib.md5(text.encode()).hexdigest()
            
            if text_hash in self._embedding_cache:
                # Use cached embedding
                self.documents.append(doc)
                self.vectors.append(self._embedding_cache[text_hash])
            else:
                if text not in text_to_indices:
                    text_to_indices[text] = []
                text_to_indices[text].append(i)
                texts.append(text)
        
        # Batch embed new texts (Together.AI supports batch embedding)
        if texts:
            try:
                # Process in smaller batches if needed
                batch_size = 25
                all_embeddings = []
                
                for i in range(0, len(texts), batch_size):
                    batch = texts[i:i+batch_size]
                    try:
                        embeddings = await asyncio.get_event_loop().run_in_executor(
                            None, self.embeddings.embed_documents, batch
                        )
                        all_embeddings.extend(embeddings)
                    except Exception as e:
                        # Fallback to individual embedding
                        logger.warning(f"Batch embedding failed, using individual: {e}")
                        for text in batch:
                            try:
                                emb = await asyncio.get_event_loop().run_in_executor(
                                    None, self.embeddings.embed_query, text
                                )
                                all_embeddings.append(emb)
                            except:
                                all_embeddings.append(None)
                
                # Store results
                for text, embedding in zip(texts, all_embeddings):
                    if embedding is not None:
                        text_hash = hashlib.md5(text.encode()).hexdigest()
                        self._embedding_cache[text_hash] = embedding
                        
                        for idx in text_to_indices.get(text, []):
                            if idx < len(documents):
                                self.documents.append(documents[idx])
                                self.vectors.append(embedding)
                
            except Exception as e:
                logger.error(f"Embedding error: {e}")
        
        embedding_time = time.time() - start_time
        logger.info(f"ULTRA SPEED: Embedded in {embedding_time:.1f}s")
    
    def similarity_search_fast(self, query_vector: np.ndarray, k: int = 5) -> List[Document]:
        """Ultra fast similarity search using pre-computed query vector"""
        if not self.vectors:
            return []
        
        try:
            # Convert to numpy for faster computation
            vectors_array = np.array(self.vectors)
            query_array = np.array(query_vector)
            
            # Fast cosine similarity (skip normalization for speed)
            similarities = np.dot(vectors_array, query_array)
            top_indices = np.argpartition(similarities, -k)[-k:]
            top_indices = top_indices[np.argsort(similarities[top_indices])][::-1]
            
            return [self.documents[i] for i in top_indices if i < len(self.documents)]
            
        except Exception as e:
            logger.error(f"Search error: {e}")
            return self.documents[:k] if len(self.documents) >= k else self.documents

async def ultra_speed_document_loader(url: str) -> List[Document]:
    """Ultra speed async document loader"""
    try:
        # Use aiohttp for async loading
        timeout = aiohttp.ClientTimeout(total=8)  # Strict timeout
        async with aiohttp.ClientSession(timeout=timeout) as session:
            async with session.get(url, headers={'User-Agent': 'Mozilla/5.0'}) as response:
                if response.status != 200:
                    raise HTTPException(status_code=400, detail=f"Failed to fetch document: {response.status}")
                
                content = await response.read()
                content_type = response.headers.get('content-type', '').lower()
        
        url_lower = url.lower()
        
        if url_lower.endswith('.pdf') or 'pdf' in content_type:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            total_pages = len(pdf_reader.pages)
            
            # ULTRA SPEED: Process even fewer pages
            if total_pages <= 15:
                pages_to_process = list(range(total_pages))
            else:
                # Minimal sampling
                first_pages = list(range(8))  # First 8 pages
                last_pages = list(range(total_pages-5, total_pages))  # Last 5 pages
                pages_to_process = first_pages + last_pages
            
            text_parts = []
            for i in pages_to_process:
                if i < len(pdf_reader.pages):
                    page_text = pdf_reader.pages[i].extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            logger.info(f"ULTRA SPEED: Processed {len(pages_to_process)}/{total_pages} PDF pages")
            return [Document(page_content=text.strip()[:60000], metadata={"source": url, "type": "pdf"})]
        
        elif url_lower.endswith('.docx') or 'wordprocessingml' in content_type:
            docx_file = io.BytesIO(content)
            doc = DocxDocument(docx_file)
            # Limit paragraphs for speed
            text = "\n".join([p.text for p in doc.paragraphs[:500] if p.text.strip()])
            return [Document(page_content=text.strip()[:60000], metadata={"source": url, "type": "docx"})]
        
        else:
            soup = BeautifulSoup(content, 'html.parser', features='lxml')  # lxml is faster
            for script in soup(["script", "style", "nav", "footer", "header", "aside"]):
                script.decompose()
            text = soup.get_text(separator=' ', strip=True)
            return [Document(page_content=text[:60000], metadata={"source": url, "type": "html"})]
        
    except asyncio.TimeoutError:
        raise HTTPException(status_code=408, detail="Document fetch timeout")
    except Exception as e:
        logger.error(f"Document load error: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to load document: {str(e)}")

class UltraSpeedRAGEngine:
    def __init__(self):
        self.chat_model = None
        self.embeddings = None
        self.text_splitter = None
        self.initialized = False
        self.vectorstore_cache = {}
        self.query_cache = {}  # Cache for questions
        self._lock = asyncio.Lock()
    
    def initialize(self):
        if self.initialized:
            return
            
        logger.info("Initializing ULTRA SPEED RAG engine...")
        
        try:
            api_key = os.getenv("TOGETHER_API_KEY", "deb14836869b48e01e1853f49381b9eb7885e231ead3bc4f6bbb4a5fc4570b78")
            os.environ["TOGETHER_API_KEY"] = api_key
            
            # Together.AI embeddings with caching
            self.embeddings = TogetherEmbeddings(
                model="BAAI/bge-base-en-v1.5",
                model_kwargs={"timeout": 5}  # Add timeout
            )
            
            # Faster model with lower latency
            self.chat_model = ChatTogether(
                model="meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
                temperature=0,
                max_tokens=2000,  # Further reduced
                timeout=10  # Add timeout
            )

            # Optimized chunking
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=800,    # Smaller chunks
                chunk_overlap=50,  # Minimal overlap
                separators=["\n\n", "\n", ". "]  # Fewer separators
            )

            self.initialized = True
            logger.info("ULTRA SPEED RAG engine ready!")
            
        except Exception as e:
            logger.error(f"Initialization error: {e}")
            raise

    async def _ultra_speed_query(self, vectorstore: UltraSpeedVectorStore, questions: List[str]) -> List[str]:
        """Ultra speed batch query with caching"""
        # Check cache first
        cache_key = hashlib.md5("|".join(questions).encode()).hexdigest()
        if cache_key in self.query_cache:
            logger.info("CACHE HIT for questions!")
            return self.query_cache[cache_key]
        
        # Embed all questions in parallel
        query_embeddings = await asyncio.gather(*[
            asyncio.get_event_loop().run_in_executor(None, self.embeddings.embed_query, q)
            for q in questions
        ])
        
        # Get context for all questions at once
        all_docs = []
        for query_emb in query_embeddings:
            docs = vectorstore.similarity_search_fast(query_emb, k=4)  # Fewer docs
            all_docs.extend(docs)
        
        # Deduplicate and limit context
        seen = set()
        unique_docs = []
        for doc in all_docs:
            doc_hash = hashlib.md5(doc.page_content.encode()).hexdigest()
            if doc_hash not in seen:
                seen.add(doc_hash)
                unique_docs.append(doc)
                if len(unique_docs) >= 8:  # Limit total docs
                    break
        
        context = " ".join([doc.page_content for doc in unique_docs])[:2500]  # Shorter context
        
        from langchain_core.messages import HumanMessage, SystemMessage
        
        # Ultra-concise prompt
        system_content = "Insurance expert. Answer questions separated by '|' with answers separated by '|'. Be accurate and specific."
        
        questions_str = " | ".join(questions)
        human_content = f"Questions: {questions_str}\n\nContext: {context}\n\nAnswer each with '|' separator:"
        
        messages = [
            SystemMessage(content=system_content),
            HumanMessage(content=human_content)
        ]
        
        try:
            response = await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(
                    None, self.chat_model.invoke, messages
                ),
                timeout=8.0  # Strict LLM timeout
            )
            
            # Parse answers
            answers = [ans.strip() for ans in response.content.split("|")]
            
            # Ensure correct number of answers
            while len(answers) < len(questions):
                answers.append("Information not found.")
            answers = answers[:len(questions)]
            
            # Cache the result
            self.query_cache[cache_key] = answers
            
            return answers
            
        except asyncio.TimeoutError:
            logger.warning("LLM timeout - using fallback")
            return ["Processing timeout - please try again." for _ in questions]

    async def process_ultra_speed(self, url: str, questions: List[str]) -> List[str]:
        """ULTRA SPEED processing with 25-second internal timeout"""
        if not self.initialized:
            raise RuntimeError("RAG engine not initialized")
        
        try:
            # Use 25 seconds to leave buffer for response
            return await asyncio.wait_for(
                self._process_internal_ultra(url, questions),
                timeout=25.0
            )
        except asyncio.TimeoutError:
            logger.error("25 second internal timeout hit")
            # Return partial results if possible
            return ["Request timeout - please try with fewer questions." for _ in questions]

    async def _process_internal_ultra(self, url: str, questions: List[str]) -> List[str]:
        url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
        
        async with self._lock:  # Prevent duplicate processing
            if url_hash in self.vectorstore_cache:
                vectorstore = self.vectorstore_cache[url_hash]
                logger.info("VECTORSTORE CACHE HIT!")
            else:
                # Load document with timeout
                docs = await ultra_speed_document_loader(url)
                
                # Split with strict limits
                chunks = self.text_splitter.split_documents(docs)
                
                # AGGRESSIVE limiting for speed
                if len(chunks) > 40:
                    # Smart sampling: beginning, key sections, end
                    chunks = chunks[:25] + chunks[-15:]
                    logger.info(f"ULTRA SPEED: Limited to {len(chunks)} chunks")
                
                vectorstore = UltraSpeedVectorStore(self.embeddings)
                await vectorstore.add_documents_ultra_fast(chunks)
                
                # Cache aggressively (keep only last 10 to save memory)
                if len(self.vectorstore_cache) >= 10:
                    # Remove oldest entry
                    oldest = list(self.vectorstore_cache.keys())[0]
                    del self.vectorstore_cache[oldest]
                
                self.vectorstore_cache[url_hash] = vectorstore
        
        # Process questions
        query_start = time.time()
        answers = await self._ultra_speed_query(vectorstore, questions)
        query_time = time.time() - query_start
        
        logger.info(f"LLM completed in {query_time:.1f}s")
        
        return answers

# Global engine
rag_engine = UltraSpeedRAGEngine()

def verify_token(authorization: Optional[str] = Header(None)):
    if authorization is None or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Authorization required")
    
    token = authorization.split("Bearer ")[-1]
    if token != EXPECTED_TOKEN:
        raise HTTPException(status_code=403, detail="Invalid token")

@asynccontextmanager
async def lifespan(app: FastAPI):
    try:
        rag_engine.initialize()
        logger.info("ULTRA SPEED RAG ready - 30s guarantee")
    except Exception as e:
        logger.error(f"Startup error: {e}")
    yield

app = FastAPI(title="ULTRA SPEED RAG API", version="4.0.0", lifespan=lifespan)

@app.post("/hackrx/run", response_model=AnswerResponse)
async def ask_questions(
    request: QuestionRequest,
    authorization: str = Depends(verify_token)
):
    """ULTRA SPEED processing - 30 second guarantee"""
    try:
        total_start = time.time()
        logger.info(f"ULTRA SPEED: {len(request.questions)} questions")

        if not request.documents.startswith(('http://', 'https://')):
            raise HTTPException(status_code=400, detail="Invalid document URL")
        if not request.questions:
            raise HTTPException(status_code=400, detail="No questions provided")
        if len(request.questions) > 20:
            raise HTTPException(status_code=400, detail="Maximum 20 questions allowed")

        answers = await rag_engine.process_ultra_speed(request.documents, request.questions)
        
        total_time = time.time() - total_start
        logger.info(f"ULTRA SPEED: Completed in {total_time:.1f}s (Target: <30s)")
        
        return {"answers": answers}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Processing error: {e}")
        # Return graceful degradation
        return {"answers": ["Error processing request. Please try again." for _ in request.questions]}

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "cache_entries": len(rag_engine.vectorstore_cache),
        "query_cache_entries": len(rag_engine.query_cache),
        "mode": "ultra_speed_optimized",
        "target_time": "30_seconds",
        "version": "4.0.0"
    }

@app.get("/")
async def root():
    return {"message": "ULTRA SPEED RAG API - 30 Second Guarantee"}

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
